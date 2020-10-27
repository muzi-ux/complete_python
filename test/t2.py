#! usr/bin/python
# -*- coding:utf-8 -*-
"""
测试对word文档的操作
"""
import os
import docx
from docx import Document
import pypinyin
import re


def pinyin(word):
    s = ''
    for i in pypinyin.pinyin(word, style=pypinyin.NORMAL):
        s += ''.join(i)
    return s


def tq(filePath, mychdir):
    lis = []
    lis2 = []
    jg_list = []
    os.chdir(mychdir)
    try:
        for i, a, k in os.walk(filePath):
            for d in k:
                one = {}
                yin = pinyin(d)
                p = re.compile(r'.docx')
                yin = p.sub("", yin)
                doc = Document(d)
                paragraph = doc.paragraphs
                for paragraph1 in doc.paragraphs:
                    lis.append(paragraph1.text)
                # 合并列表lis2，添加至字典，字典名是文件的拼音
                lis2 = ["".join(x for x in lis)]
                # one[yin] = lis2
                jg_list.append(lis2)
                lis2 = []
    except Ellipsis:
        pass
    return jg_list


def txt(lis):
    do = Document()
    for i in lis:
        for a in i:
            do.add_paragraph(str(a))
        do.save("Python_word.docx")
    print("完成")


# os.chdir()
# hz = Document("汇总.docx")
# hz.add_paragraph('Add new paragraph', style='ListNumber')

# doc = Document()
# paragraph = doc.paragraphs[1]
# runs = paragraph.runs
# print(runs)
# for run in runs:
#     print(run.text)


if __name__ == '__main__':
    tq()
